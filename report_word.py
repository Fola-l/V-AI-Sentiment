"""
Generate an editable Word (.docx) report from analysis outputs.
Mirrors the structure of the PDF but fully editable in Word/Google Docs.

Usage:
    python report_word.py
"""

from pathlib import Path
from datetime import datetime

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── colour helpers ────────────────────────────────────────────────────────────

def rgb(r, g, b):
    return RGBColor(r, g, b)

GREEN  = rgb(76, 175, 80)
RED    = rgb(244, 67, 54)
GREY   = rgb(120, 120, 120)
DARK   = rgb(30, 30, 30)
NAVY   = rgb(25, 25, 60)
LIGHT  = rgb(240, 240, 240)

def set_cell_bg(cell, r, g, b):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)


# ── load data ─────────────────────────────────────────────────────────────────

def load_data():
    df        = pd.read_csv("processed_tweets.csv", dtype=str)
    rc_themes = pd.read_csv("root_cause_themes.csv")
    rc_quotes = pd.read_csv("root_cause_quotes.csv", dtype=str)
    rc_terms  = pd.read_csv("root_cause_summary.csv", dtype=str)

    df["compound"] = pd.to_numeric(df["compound"], errors="coerce").fillna(0)
    df["clean_text"] = df["clean_text"].fillna("")
    df["themes_str"] = df["themes_str"].fillna("untagged")

    return df, rc_themes, rc_quotes, rc_terms


def sanitize(text: str) -> str:
    replacements = {"—": "-", "–": "-", "−": "-", "…": "...", "’": "'", "‘": "'"}
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


# ── document helpers ──────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, colour=DARK):
    p = doc.add_heading(sanitize(text), level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = colour
    return p


def add_body(doc, text, italic=False):
    p = doc.add_paragraph(sanitize(text))
    p.style.font.size = Pt(11)
    if italic:
        for run in p.runs:
            run.italic = True
    return p


def add_kv(doc, key, value, colour=None):
    p = doc.add_paragraph()
    k = p.add_run(f"{key}  ")
    k.bold = True
    k.font.color.rgb = GREY
    v = p.add_run(sanitize(str(value)))
    v.bold = False
    if colour:
        v.font.color.rgb = colour


def add_bullet(doc, text):
    p = doc.add_paragraph(sanitize(text), style="List Bullet")
    p.style.font.size = Pt(11)


def add_quote(doc, text, sentiment, compound):
    colour = GREEN if sentiment == "positive" else RED if sentiment == "negative" else GREY
    p = doc.add_paragraph()
    run = p.add_run(f'"{sanitize(text[:300])}"')
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    meta = doc.add_paragraph()
    m = meta.add_run(f"  compound: {float(compound):.3f} | sentiment: {sentiment}")
    m.font.size = Pt(9)
    m.font.color.rgb = colour
    doc.add_paragraph()  # spacer


def add_image(doc, path, width=Inches(6)):
    if path and Path(path).exists():
        doc.add_picture(str(path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def divider(doc):
    doc.add_paragraph("_" * 80)


# ── build ─────────────────────────────────────────────────────────────────────

def build_word(df, rc_themes, rc_quotes, rc_terms):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    total   = len(df)
    n_pos   = int((df["sentiment"] == "positive").sum())
    n_neg   = int((df["sentiment"] == "negative").sum())
    n_neu   = int((df["sentiment"] == "neutral").sum())
    pct_pos = n_pos / total * 100
    pct_neg = n_neg / total * 100
    pct_neu = n_neu / total * 100
    avg_c   = df["compound"].mean()

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading("UK Twitter / X Sentiment Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph("AI Voice Agents Answering Service Calls")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = GREY

    date_p = doc.add_paragraph(f"Generated {datetime.now().strftime('%d %B %Y')}  |  Data: twitterapi.io  |  NLP: VADER")
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in date_p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = GREY

    doc.add_paragraph()

    # ── Key stats table ───────────────────────────────────────────────────────
    stats_table = doc.add_table(rows=1, cols=4)
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats_table.style = "Table Grid"
    cells = stats_table.rows[0].cells
    for cell, val, label, r, g, b in [
        (cells[0], str(total),        "Tweets analysed", 60, 60, 80),
        (cells[1], f"{pct_pos:.0f}%", "Positive",        76, 175, 80),
        (cells[2], f"{pct_neu:.0f}%", "Neutral",         120, 120, 130),
        (cells[3], f"{pct_neg:.0f}%", "Negative",        244, 67, 54),
    ]:
        set_cell_bg(cell, r, g, b)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_val = p.add_run(val + "\n")
        run_val.bold = True
        run_val.font.size = Pt(20)
        run_val.font.color.rgb = RGBColor(255, 255, 255)
        run_lbl = p.add_run(label)
        run_lbl.font.size = Pt(9)
        run_lbl.font.color.rgb = RGBColor(220, 220, 220)

    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    add_heading(doc, "1. Executive Summary", 1)
    add_body(doc,
        f"This report analyses {total} UK-relevant tweets collected via the Twitter / X API "
        "across 19 search queries covering AI voice agents, automated phone systems, virtual "
        "receptionists, and related topics. Sentiment was scored using the VADER model "
        "(compound range -1 to +1) and tweets were tagged against 10 thematic categories."
    )

    add_heading(doc, "Key Numbers", 2)
    add_kv(doc, "Total tweets:", total)
    add_kv(doc, "Positive:", f"{n_pos}  ({pct_pos:.1f}%)", GREEN)
    add_kv(doc, "Neutral:",  f"{n_neu}  ({pct_neu:.1f}%)", GREY)
    add_kv(doc, "Negative:", f"{n_neg}  ({pct_neg:.1f}%)", RED)
    add_kv(doc, "Avg compound score:", f"{avg_c:+.3f}")

    add_heading(doc, "Top-Line Findings", 2)
    bullets = [
        "Reactions are broadly mixed-to-positive overall (~47% positive), but highly context-dependent.",
        "Low-stakes use cases (restaurant bookings, taxi rides, appointment reminders) attract the most positive responses.",
        "Negative sentiment clusters around four triggers: deception, blocked human access, accuracy failures, and high-stakes contexts (NHS, GP, banking).",
        "Transparency is the single largest lever - upfront AI disclosure skews reactions strongly positive.",
        "Resistance in healthcare and banking is categorical (a principle), not experiential - design improvements alone will not resolve it.",
    ]
    for b in bullets:
        add_bullet(doc, b)

    # ── Sentiment Distribution ────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "2. Sentiment Distribution", 1)

    add_image(doc, "charts/sentiment_pie.png", width=Inches(3.5))
    add_image(doc, "charts/compound_histogram.png", width=Inches(5.5))

    add_heading(doc, "Compound Score Interpretation", 2)
    add_body(doc,
        "VADER compound scores range from -1.0 (maximally negative) to +1.0 (maximally positive). "
        "Tweets are classified positive (>= 0.05), negative (<= -0.05), or neutral (between). "
        "The histogram shows a bimodal distribution - strong peaks at both ends confirm polarised "
        "opinion rather than a simple majority view."
    )

    if Path("charts/timeline.png").exists():
        add_heading(doc, "Tweet Volume Over Time", 2)
        add_image(doc, "charts/timeline.png", width=Inches(6))

    # ── Theme Analysis ────────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "3. Theme Analysis", 1)

    add_heading(doc, "Sentiment Split by Theme", 2)
    add_image(doc, "charts/theme_sentiment_stacked.png", width=Inches(6))

    add_heading(doc, "Average Compound Score by Theme", 2)
    add_image(doc, "charts/avg_compound_theme.png", width=Inches(6))

    add_heading(doc, "Theme x Sentiment Heatmap", 2)
    add_image(doc, "charts/theme_heatmap.png", width=Inches(5.5))

    # Theme table
    add_heading(doc, "Theme Polarity Table", 2)
    headers = ["Theme", "Count", "Pos %", "Neu %", "Neg %", "Avg Score", "Dominant"]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr_cells[i], 60, 60, 80)
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9)

    for _, row in rc_themes.sort_values("avg_compound", ascending=False).iterrows():
        dom = str(row.get("dominant_sentiment", ""))
        row_cells = t.add_row().cells
        data = [
            str(row["theme"]),
            str(int(row["count"])),
            f"{row['pct_positive']:.1f}",
            f"{row['pct_neutral']:.1f}",
            f"{row['pct_negative']:.1f}",
            f"{row['avg_compound']:+.3f}",
            dom,
        ]
        col_colour = GREEN if dom == "positive" else RED if dom == "negative" else GREY
        for i, val in enumerate(data):
            p = row_cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if i in (0, 6):
                run.font.color.rgb = col_colour

    doc.add_paragraph()

    # ── Root Cause: Language Drivers ──────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "4. Root Cause - Language Drivers", 1)

    add_heading(doc, "What Drives Positive Sentiment", 2)
    pos_terms = rc_terms["positive_term"].dropna().head(12).tolist()
    add_body(doc, "Distinctive language in positive tweets:")
    add_body(doc, "  " + "  |  ".join(pos_terms), italic=True)
    add_body(doc,
        "Positive language clusters around efficiency and convenience - words like 'quick', "
        "'easy', 'saves time' suggest people appreciate AI voice agents when they make a task "
        "faster, particularly for low-stakes interactions (bookings, confirmations, basic queries)."
    )
    add_image(doc, "charts/tfidf_positive.png", width=Inches(6))

    add_heading(doc, "What Drives Negative Sentiment", 2)
    neg_terms = rc_terms["negative_term"].dropna().head(12).tolist()
    add_body(doc, "Distinctive language in negative tweets:")
    add_body(doc, "  " + "  |  ".join(neg_terms), italic=True)
    add_body(doc,
        "Negative language centres on frustration ('useless', 'waste of time'), deception "
        "('pretending', 'didn't say it was AI'), and access barriers ('can't get through', "
        "'just want to speak to someone')."
    )
    add_image(doc, "charts/tfidf_negative.png", width=Inches(6))

    # ── Representative Quotes ─────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "5. Representative Quotes (Anonymised)", 1)

    add_heading(doc, "Strongest Positive Reactions", 2)
    for _, r in rc_quotes[rc_quotes["sentiment"] == "positive"].head(5).iterrows():
        add_quote(doc, str(r["clean_text"]), "positive", r["compound"])

    add_heading(doc, "Strongest Negative Reactions", 2)
    for _, r in rc_quotes[rc_quotes["sentiment"] == "negative"].head(5).iterrows():
        add_quote(doc, str(r["clean_text"]), "negative", r["compound"])

    # ── Research Questions ────────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "6. Research Questions Answered", 1)

    qa = [
        ("Q1 - Are reactions mostly positive, neutral, or negative?",
         f"Broadly mixed-positive: {pct_pos:.0f}% positive, {pct_neu:.0f}% neutral, "
         f"{pct_neg:.0f}% negative. Sentiment bifurcates sharply by use-case context."),
        ("Q2 - What are the strongest concerns?",
         "In order: (1) Deception - AI not identifying itself. (2) No human escape route. "
         "(3) Accuracy / mishearing errors. (4) Inappropriate use in high-stakes domains "
         "(NHS, GP, banking). (5) Privacy and call recording."),
        ("Q3 - What conditions make AI voice agents more acceptable?",
         "Low-stakes task + transparent disclosure + fast human fallback + accurate "
         "comprehension. All four conditions together produce consistently positive sentiment."),
        ("Q4 - Does a natural/human-like voice seem helpful, creepy, or both?",
         "Both. A smooth voice reduces friction for routine tasks but triggers uncanny valley "
         "concerns when users realise mid-call it's AI - especially in emotionally sensitive contexts."),
        ("Q5 - What design recommendations follow?",
         "See next section."),
    ]
    for q, a in qa:
        add_heading(doc, q, 2)
        add_body(doc, a)

    # ── Design Recommendations ────────────────────────────────────────────────
    doc.add_page_break()
    add_heading(doc, "7. Design Recommendations", 1)

    recs = [
        ("1. Disclose upfront - always",
         "Open with 'Hi, I'm an AI assistant for [company].' Deception is the single "
         "largest driver of negative sentiment."),
        ("2. Instant, reliable human escape hatch",
         "'Say agent at any time' must be prominent, fast, and actually work."),
        ("3. Confine AI to low-stakes flows",
         "Bookings, confirmations, FAQs, status updates. Route NHS, GP, banking, "
         "complaints, and emergencies to humans."),
        ("4. Fix accent and dialect comprehension",
         "Scottish, Welsh, Geordie, and other regional UK accents are a consistent "
         "pain point. Test with non-RP speakers before deployment."),
        ("5. State data handling proactively",
         "Mention call recording policy at the start of the call."),
        ("6. Make errors recoverable instantly",
         "Misheard inputs should trigger an immediate clarification loop, not a dead end."),
    ]
    for title, body in recs:
        add_heading(doc, title, 2)
        add_body(doc, body)

    divider(doc)
    add_heading(doc, "Conclusion", 1)
    add_body(doc,
        "People are not automatically against AI answering calls. Acceptance is higher for "
        "simple, low-stakes interactions when the AI is transparent, accurate, and offers a "
        "fast path to a human. The biggest design failure is deception; the biggest commercial "
        "opportunity is frictionless task completion for routine bookings."
    )

    out = "AI_Voice_Agent_Report.docx"
    doc.save(out)
    print(f"Word report saved to {out}")


def main():
    print("Loading data ...")
    df, rc_themes, rc_quotes, rc_terms = load_data()
    print("Building Word report ...")
    build_word(df, rc_themes, rc_quotes, rc_terms)


if __name__ == "__main__":
    main()
